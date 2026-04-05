from datetime import datetime
from io import BytesIO, StringIO
import docx, csv, json, requests, threading, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from PIL import Image, ImageDraw, ImageFont

from app.services.register import get_active_registers
from app.services.room import get_all_rooms
from app.services.transaction import get_all_transaction_api

DEFAULT_CONFIG = {
    "DISCORD_WEBHOOK_URL": "",
    "DISCORD_BOT_TOKEN": "",
    "HOTEL_NAME": "",
    "HOTEL_ADDRESS": "",
    "REPORT_PATH": ".",
    "MONTHLY_REPORT_FOLDER": "", # 'True' for a success value
    "DISCORD_CHANNEL_ID_UPDATES": "",
    "DISCORD_CHANNEL_ID_TEST": ""
}

def get_config():
    return json.load(open('instance\\config.json', 'r'))

def ct(t):
    days = t.days
    hours = t.seconds // 3600
    return f'{days} Day {hours} Hour'

def format_dt(dt):
    return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else None

def calculate_total_balance(rent_per_day, start_time, end_time = None):
    days, hours = calculate_time_difference(start_time, end_time, return_str=False)
    if days == 0:
        days = 1
    elif hours > 1.8: # type: ignore
        days += 1 # type: ignore
    return rent_per_day * days

def calculate_time_difference(start_time, end_time = None, return_str = False):
    if end_time == None:
        end_time = datetime.now()
    diff = end_time - start_time
    if return_str:
        return ct(diff)
    return diff.days, diff.seconds / 3600

def model_to_csv(model, data):
    output = StringIO()
    writer = csv.writer(output)

    columns = [col.name for col in model.__table__.columns]
    writer.writerow(columns)

    for row in data:
        writer.writerow([getattr(row, col) for col in columns])

    return output.getvalue()


def register_records_to_array(entries):
    register = [
        {
            'reg_id': (register.reg_id) if register.reg_id else ('C' + str(register.id)),
            'id': register.id,
            'room_number': register.room.room_number,
            'name': register.customer.name,
            'address': register.customer.address,
            'phone': register.customer.phone,
            'email': register.customer.email,
            'id_type': register.customer.id_type,
            'id_detail': register.customer.id_detail,
            'check_in': format_dt(register.check_in),
            'check_out': format_dt(register.check_out),
            'pov': register.purpose_of_visit,
            'rent_per_day': register.rent_per_day,
            'time_passed': ct((register.check_out or datetime.now()) - register.check_in),
            'total_paid': register.total_paid,
            'total_balance': calculate_total_balance(register.rent_per_day, register.check_in, register.check_out),
            'remaining_balance': calculate_total_balance(register.rent_per_day, register.check_in, register.check_out) - register.total_paid,
        } for register in entries
    ]
    return register

def transactions_to_array(records):
    n_records = []
    for record in records:
        if record:
            n_records.append({
                'name': record.name,
                'transaction_time': record.transaction_time.strftime('%d-%m-%y %H:%M'),
                'payment_mode': record.payment_method,
                'room_no': record.room_number,
                'amount_paid': record.amount_paid
            })
    return n_records

def generate_docx_report(entries, date):
    HOTEL_NAME = get_config().get('HOTEL_NAME', '')
    HOTEL_ADDRESS = get_config().get('HOTEL_ADDRESS', '')
    buffer = BytesIO()

    entries = register_records_to_array(entries)

    doc = docx.Document()
    # para_1 = doc.add_paragraph('Mobile: ')
    # para_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header1 = doc.add_heading(HOTEL_NAME)
    header1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.style.font.size = Pt(36) # type: ignore
    header2 = doc.add_heading(HOTEL_ADDRESS, level=3)
    header2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.style.font.size = Pt(24) # type: ignore

    para_date = doc.add_paragraph('Date: ' + date)
    para_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = doc.add_table(1, 6, style='Table Grid')
    headers = ['Sl. No.', 'Name\nAddress', 'Phone', 'Check In', 'Document', 'Purpose Of Visit']
    col_widths = [Inches(0.7), Inches(2), Inches(1), Inches(1), Inches(1.2), Inches(1)]
    for idx in range(len(headers)):
        table.rows[0].cells[idx].text = headers[idx]
    
    for idx, record in enumerate(entries):
        row = table.add_row()

        row.cells[0].text = str(idx + 1)
        row.cells[1].text = record['name'] + '\n' + record['address']
        row.cells[2].text = str(record['phone'])
        row.cells[3].text = record['check_in'].split(' ')[1]
        row.cells[4].text = record['id_type'] + '\n' + record['id_detail']
        row.cells[5].text = record['pov']
    table.autofit = False

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
    
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def _send_webhook_alert(data):
    config = get_config()
    DISCORD_WEBHOOK_URL = config['DISCORD_WEBHOOK_URL']
    try:
        res = requests.post(DISCORD_WEBHOOK_URL, json={'embeds': data})
        res.raise_for_status()
        return {'success': True, 'message': 'Alerted Successfully'}
    except requests.exceptions.RequestException as err:
        file_path = 'instance/notify.json'
        if not os.path.exists(file_path):
            with open(file_path, 'w') as f:
                json.dump([], f)
        try:
            with open(file_path, 'r') as f:
                load_data = json.load(f)
        except json.JSONDecodeError:
            load_data = []
        for i in data:
            load_data.append(i)
        with open(file_path, 'w') as f:
            json.dump(load_data, f, indent=4)

        return {'success': False, 'message': 'Saved for retry'}

def send_webhook_alert(data):
    threading.Thread(None, _send_webhook_alert, args=([data], )).start()

def push_notification():
    file_path = 'instance/notify.json'
    if not os.path.exists(file_path):
            with open(file_path, 'w') as f:
                json.dump([], f)
    try:
        with open(file_path, 'r') as f:
            load_data = json.load(f)
    except json.JSONDecodeError:
        load_data = []
    with open(file_path, 'w') as f:
        json.dump([], f)
    
    for i in load_data:
        _send_webhook_alert(i)

def generate_image(rooms = None, registers = None, transactions = None):
    if rooms is None:
        rooms = get_all_rooms()
    if registers is None:
        registers = get_active_registers()
    if transactions is None:
        transactions = get_all_transaction_api(transaction_date = datetime.now().strftime('%Y-%m-%d'))
    # Draw grid
    x, y = 20, 20
    box_w, box_h = 300, 400
    gap_x, gap_y = 40, 40
    cols = 8
    total_rows = 0
    n = len(rooms)

    # Group rooms by floor using a normal dict
    floors = {}
    for room in rooms:
        if room.floor_number not in floors:
            floors[room.floor_number] = []
        floors[room.floor_number].append(room)
    for k, v in floors.items():
        if len(v) > cols:
            total_rows += (len(v) // cols)
        else:
            total_rows += 1
    
    # Load font's
    try:
        font = ImageFont.truetype("arial.ttf", 64)
        font_small = ImageFont.truetype("arial.ttf", 32)
    except:
        font = ImageFont.load_default(64)
        font_small = ImageFont.load_default(32)
    
    width = x + cols * box_w + (cols - 1) * gap_x + x
    height = y + total_rows * box_h + (total_rows - 1) * gap_y + y + (64 * 4)
    img = Image.new("RGB", (width, height), "black")
    draw = ImageDraw.Draw(img)


    # Map room → customer
    room_map = {}
    for r in registers:
        if r is not None:
            room_map[r.room.room_number] = {
                'name': r.customer.name or 'Unknown',
                'time_passed': calculate_time_difference(r.check_in, return_str=True),
                'total_paid': r.total_paid
            }
    floor_idx = 0

    for _, (floor, floor_rooms) in enumerate(floors.items()):
        try:
            floor_rooms = sorted(floor_rooms, key=lambda r: int(r.room_number))
        except Exception as err:
            floor_rooms = sorted(floor_rooms, key=lambda r: r.room_number)

        row = 0

        for i, room in enumerate(floor_rooms):
            col = i % cols
            row = floor_idx + (i // cols)

            rx = x + col * (box_w + gap_x)
            ry = y + row * (box_h + gap_y)

            is_available = room.is_available
            color = (0, 180, 0) if is_available else (200, 0, 0)

            draw.rectangle([rx, ry, rx + box_w, ry + box_h], fill=color)

            text = f"{room.room_number}"
            draw.text((rx + 90, ry + 10), text, fill="white", font=font)
            if not is_available:
                text = f"\n{room_map.get(room.room_number, {}).get('name', '')}".replace(' ', '\n')
                text += f'\n\n{room_map.get(room.room_number, {}).get('time_passed', '')}'
                text += f'\n\nTotal Paid: {room_map.get(room.room_number, {}).get('total_paid', '')}'
                draw.text((rx+30, ry + 60), text, fill="white", font=font_small)
        
        floor_idx = row + 1
    # ---------------------------
    # 💰 TRANSACTION SUMMARY
    # ---------------------------
    upi, cash = 0, 0

    for t in transactions:
        method = t.payment_method.lower()
        if method == "upi":
            upi += t.amount_paid
        elif method == "cash":
            cash += t.amount_paid

    total = upi + cash

    summary = (
        f"UPI: ₹{upi}\n"
        f"Cash: ₹{cash}\n"
        f"Total: ₹{total}"
    )

    draw.text((20, height - 200), summary, fill="white", font=font)

    buffer = BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    # files = {
    #     'file': ('image.png', buffer, 'image/png')
    # }
    # res = requests.post(get_config()['DISCORD_WEBHOOK_URL'], files=files)
    return buffer