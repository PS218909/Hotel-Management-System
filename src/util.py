import pandas as pd
import os, datetime, requests, json, time, threading
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from .config import *

def default_values():
    if not os.path.exists(DATASET_DIR):
        os.mkdir(DATASET_DIR)
    if not os.path.exists(ROOMS_DB):
        with open(ROOMS_DB, "w") as writer:
            writer.write("r,f,s")
            for room_no in range(1, 17):
                writer.write('\n' + str(room_no) + ',0,1')
            for room_no in range(101, 105):
                writer.write('\n' + str(room_no) + ',1,1')
            for room_no in range(201, 209):
                writer.write('\n' + str(room_no) + ',2,1')
            for room_no in range(301, 309):
                writer.write('\n' + str(room_no) + ',3,1')
            for room_no in range(401, 409):
                writer.write('\n' + str(room_no) + ',4,1')
    else:
        df = read_csv(ROOMS_DB)
        cols = df.columns
        if 'r' not in cols:
            if 'room' in cols:
                df['r'] = df['room']
                df.drop('room', inplace=True)
            else:
                df['r'] = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '101', '102', '103', '104', '201', '202', '203', '204', '205', '206', '207', '208', '301', '302', '303', '304', '305', '306', '307', '308', '401', '402', '403', '404', '405', '406', '407', '408']
        if 'f' not in cols:
            if 'floor' in cols:
                df['f'] = df['floor']
                df.drop('floor', inplace=True)
            else:
                df['f'] = '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '0', '1', '1', '1', '1', '2', '2', '2', '2', '2', '2', '2', '2', '3', '3', '3', '3', '3', '3', '3', '3', '4', '4', '4', '4', '4', '4', '4', '4'
        if 's' not in cols:
            if 'status' in cols:
                df['s'] = df['status']
                df.drop('status', inplace=True)
            else:
                df['s'] = ['1', '2', '1', '1', '1', '1', '1', '1', '1', '1', '1', '1', '1', '1', '1', '1', '2', '2', '2', '1', '2', '2', '2', '2', '2', '2', '1', '2', '2', '1', '2', '1', '2', '2', '1', '2', '1', '2', '2', '1', '1', '1', '1', '1']
        df = df[['r', 'f', 's']]
        write_csv(ROOMS_DB, df)
    if not os.path.exists(CUSTOMERS_DB):
        with open(CUSTOMERS_DB, 'w') as writer:
            writer.write('id,n,a,p,it,ip')
    else:
        df = read_csv(CUSTOMERS_DB)
        cols = df.columns
        if 'id' not in cols:
            df['id'] = range(len(df))
        if 'n' not in cols:
            if 'name' in cols:
                df['n'] = df['name']
            else:
                df['n'] = [None] * len(df)
        if 'p' not in cols:
            if 'phone' in cols:
                df['p'] = df['phone']
            else:
                df['p'] = [None] * len(df)
        if 'a' not in cols:
            if 'address' in cols:
                df['a'] = df['address']
            else:
                df['a'] = [None] * len(df)
        if 'it' not in cols:
            if 'id type' in cols:
                df['it'] = df['id type']
            else:
                df['it'] = [None] * len(df)
        if 'ip' not in cols:
            if 'id detail' in cols:
                df['ip'] = df['id detail']
            else:
                df['ip'] = [None] * len(df)
        df = df[['id', 'n', 'a', 'p', 'it', 'ip']]
        write_csv(CUSTOMERS_DB, df)
    if not os.path.exists(TRANSACTIONS_DB):
        with open(TRANSACTIONS_DB, 'w') as writer:
            writer.write('id,rid,a,t,m,d')
    else:
        df = read_csv(TRANSACTIONS_DB)
        cols = df.columns
        if 'id' not in cols:
            df['id'] = range(len(df))
        if 'rid' not in cols:
            if 'register_id' in cols:
                df['rid'] = df['register_id']
            else:
                df['rid'] = [None] * len(df)
        if 'a' not in cols:
            if 'amount' in cols:
                df['a'] = df['amount']
            else:
                df['a'] = [None] * len(df)
        if 't' not in cols:
            if 'datetime' in cols:
                df['t'] = df['datetime']
            else:
                df['t'] = [None] * len(df)
        if 'm' not in cols:
            if 'mode' in cols:
                df['m'] = df['mode']
            else:
                df['m'] = [None] * len(df)
        if 'd' not in cols:
            if 'description' in cols:
                df['d'] = df['description']
            else:
                df['d'] = [None] * len(df)
        df = df[['id', 'rid', 'a', 't', 'm', 'd']]
        write_csv(TRANSACTIONS_DB, df)
    if not os.path.exists(REGISTER_DB):
        with open(REGISTER_DB, 'w') as writer:
            writer.write('id,rno,cid,ac,rpd,pov,cin,cout')
    else:
        df = read_csv(REGISTER_DB)
        cols = df.columns
        if 'id' not in cols:
            df['id'] = range(len(df))
        if 'rno' not in cols:
            if 'room' in cols:
                df['rno'] = df['room']
            else:
                df['rno'] = [None] * len(df)
        if 'cid' not in cols:
            if 'customer_id' in cols:
                df['cid'] = df['customer_id']
            else:
                df['cid'] = [None] * len(df)
        if 'ac' not in cols:
            if 'advance' in cols:
                df['ac'] = df['advance']
            else:
                df['ac'] = [None] * len(df)
        if 'rpd' not in cols:
            if 'rate_per_day' in cols:
                df['rpd'] = df['rate_per_day']
            else:
                df['rpd'] = ['0'] * len(df)
        if 'pov' not in cols:
            if 'purpose_of_visit' in cols:
                df['pov'] = df['purpose_of_visit']
            else:
                df['pov'] = ['Official'] * len(df)
        if 'cin' not in cols:
            if 'checkin' in cols:
                df['cin'] = df['checkin']
            else:
                df['cin'] = [None] * len(df)
        if 'cout' not in cols:
            if 'checkout' in cols:
                df['cout'] = df['checkout']
            else:
                df['cout'] = [None] * len(df)
        if 'gb' not in cols:
            if 'gst_bill' in cols:
                df['gb'] = df['gst_bill']
            else:
                df['gb'] = [''] * len(df)
        df = df[['id', 'rno', 'cid', 'ac', 'rpd', 'pov', 'cin', 'cout', 'gb']]
        write_csv(REGISTER_DB, df)
    if not os.path.exists(ALERT_DB):
        with open(ALERT_DB, 'w') as writer:
            writer.write('[]')

def write_csv(filename: str, df: pd.DataFrame):
    try:
        df = df.fillna('')
        for col in df.columns:
            try:
                df[col] = df[col].astype(str)
                df[col] = df[col].str.upper()
                for i in range(3):
                    df[col] = df[col].str.replace('  ', ' ')
            except Exception as err:
                continue
        df.to_csv(filename, index=False)
    except Exception as err:
        print(err)
        return err

def read_csv(filename: str) -> pd.DataFrame:
    try:
        df = pd.read_csv(filename, index_col=False)
        return df
    except Exception as err:
        print(err)
        return pd.DataFrame({})

def get_rooms_status():
    rooms_df: pd.DataFrame = read_csv(ROOMS_DB)
    data = {}
    for idx, room in rooms_df.iterrows():
        data[room['f']] = data.get(room['f'], {})
        if room['s'] == OCCUPIED:
            data[room['f']][room['r']] = {'status': room['s'], 'details': get_room_detail(room['r'])}
        else:
            data[room['f']][room['r']] = {'status': room['s'], 'details': {}}
    return data

def get_room_detail(rno):
    register_df: pd.DataFrame = read_csv(REGISTER_DB)
    customer_df: pd.DataFrame = read_csv(CUSTOMERS_DB)

    register_df['rno'] = register_df['rno'].astype(type(rno))
    register_df['cout'] = register_df['cout'].fillna('')
    register_df['gb'] = register_df['gb'].fillna('')
    register_df = register_df[(register_df['rno'] == rno) & (register_df['cout'] == '')]
    if len(register_df) > 0:
        row = register_df.loc[register_df.index[-1]]
        customer = customer_df.loc[customer_df[customer_df['id'] == row['cid']].index[0]]
        res = {
            'status': 2,
            'roomno': rno,
            'checkin': row['cin'], 
            'ac': row['ac'],
            'pov': row['pov'],
            'rpd': row['rpd'], 
            'name': customer['n'], 
            'phone': customer['p'], 
            'address': customer['a'], 
            'id_type': customer['it'],
            'id_detail': customer['ip'],
            'time_passed': time_difference_str(row['cin']),
            'payments': fetch_payments(row['id']),
            'gst_bill': row['gb'],
        }
        res['total_payment'] = sum([payment['a_x'] for payment in res['payments']])
        return {'success': True, 'message': 'Room is occupied', 'data': res}
    else:
        cnt_time = datetime.datetime.now() # + datetime.timedelta(hours=5, minutes=30)
        res = {
            'status': 1,
            'roomno': rno,
            'checkin': cnt_time.strftime('%Y-%m-%dT%H:%M'),
        }
        return {'success': False, 'message': 'Room is available.', 'data': res}

def get_register_detail(rid = None, page=None, count=None, date=None, gst=False, cid=None):
    if page == None:
        register_df: pd.DataFrame = read_csv(REGISTER_DB)
    else:
        page = page - 1
        register_df: pd.DataFrame = read_csv(REGISTER_DB)
        register_df = register_df.sort_values(by='cin', key=lambda x: pd.to_datetime(x, format='%Y-%m-%dT%H:%M'), ascending=False)
        empty_df = register_df[register_df['cout'].isna()]
        non_empty_df = register_df[register_df['cout'].notna()]
        register_df = pd.concat([empty_df, non_empty_df])
        register_df = register_df[page * count:(page + 1) * count][::-1]
        if len(register_df) <= 0:
            return pd.DataFrame({})
    if date is not None:
        register_df = register_df[register_df['cin'].str.startswith(date)]
    if rid is not None:
        register_df = register_df[register_df['id'] == int(rid)]
    if gst:
        register_df = register_df.dropna()
    if cid is not None:
        register_df = register_df[register_df['cid'] == int(cid)]
    if len(register_df) == 0:
        return pd.DataFrame(columns=['id_x', 'rno', 'rpd', 'pov', 'n', 'a', 'p', 'it', 'ip', 'time_passed', 'amount_paid', 'cin', 'cout', 'remaining_balance'])
    customer_df: pd.DataFrame = read_csv(CUSTOMERS_DB)
    merged_register_customer = register_df.merge(right=customer_df, how='left', left_on='cid', right_on='id')
    merged_register_customer['amount_paid'] = merged_register_customer.apply(lambda row: sum([payment['a_x'] for payment in fetch_payments(row['id_x'])]), axis=1)
    merged_register_customer = merged_register_customer.sort_values(by='cin', key=lambda x: pd.to_datetime(x, format='%Y-%m-%dT%H:%M'), ascending=False)
    empty_df = merged_register_customer[merged_register_customer['cout'].isna()]
    non_empty_df = merged_register_customer[merged_register_customer['cout'].notna()]
    merged_register_customer = pd.concat([empty_df, non_empty_df])
    merged_register_customer['cout'] = merged_register_customer['cout'].fillna('')
    merged_register_customer['time_passed'] = merged_register_customer.apply(lambda row: time_difference_str(row['cin'], row['cout']), axis=1)
    merged_register_customer['remaining_balance'] = merged_register_customer.apply(lambda row: ((time_difference(row['cin'], row['cout'])[0] + (1 if time_difference(row['cin'], row['cout'])[1] >= 0 else 0))*row['rpd']) - row['amount_paid'], axis=1)
    merged_register_customer['cin'] = merged_register_customer.apply(lambda row: transform_time(row['cin']), axis=1) # type: ignore
    merged_register_customer['cout'] = merged_register_customer.apply(lambda row: transform_time(row['cout']), axis=1) # type: ignore
    final_df = merged_register_customer[['id_x', 'rno', 'rpd', 'pov', 'n', 'a', 'p', 'it', 'ip', 'time_passed', 'amount_paid', 'cin', 'cout', 'remaining_balance', 'gb']]
    final_df['cout'] = final_df['cout'].fillna('Present')
    return final_df

def fetch_gst_info():
    register_df = read_csv(REGISTER_DB)
    transaction_df = read_csv(TRANSACTIONS_DB)

    register_df = register_df.dropna(subset=['gb'])

    merged = transaction_df.merge(register_df, how='inner', left_on='rid', right_on='id')
    merged = merged.sort_values(by='gb')
    merged_required = merged[['gb', 'a', 'm', 'rno', 't']]
    merged_required['gb'] = merged_required['gb'].astype(int)
    merged_required['t'] = merged_required['t'].apply(transform_time)
    return merged_required

def append_row(file, row):
    df = read_csv(file)
    try:
        if isinstance(row, list):
            df.loc[len(df)] = row
            write_csv(file, df)
        if isinstance(row, dict):
            length = len(df)
            for key, value in row.items():
                df.loc[length, key] = value
            write_csv(file, df)
        return {'success': True, 'message': 'Data Added Successfully'}
    except Exception as err:
        return {'success': False, 'message': 'Failed to add data because' + str(err)}

def update_row(file: str, query: dict, row: dict):
    df = read_csv(file)
    df = df.fillna('')
    try:
        for rule in query:
            df[rule] = df[rule].astype(type(query[rule]))
        condition = (df[list(query)] == pd.Series(query)).all(axis=1)
        if condition.any():
            index = df[condition].index[0]
            df.loc[index, list(row.keys())] = list(row.values())
            write_csv(file, df)
        else:
            return {'success': False, 'message': 'Unable to locate the conditions'}
        return {'success': True, 'message': 'Data Updated Successfully'}
    except Exception as err:
        return {'success': False, 'message': 'Failed to update because ' + str(err)}

def delete_row(file: str, query: dict):
    res = find_by(file=file, query=query)
    df = read_csv(file)
    if len(res) > 0:
        df = df.drop(index=res.index[0])
        write_csv(TRANSACTIONS_DB, df)
        return {'success': True, 'message': 'Row deleted successfully'}
    else:
        return {'success': False, 'message': 'No such data found.'}

def find_by(query, file):
    df = read_csv(file)
    for rule in query:
        if query[rule] is None:
            df = df.loc[df[rule].isna()]
        else:
            df[rule] = df[rule].astype(type(query[rule]))
            df = df.loc[df[rule] == query[rule]]
    return df

def next_id(file, id='id'):
    df = read_csv(file)
    df[id] = df[id].astype(int)
    if len(df) == 0:
        return 1
    return df[id].max() + 1

def search_customer(name, phone):
    df = read_csv(CUSTOMERS_DB)
    return df[(df['n'].str.lower().str.startswith(name.lower())) & (df['p'].astype(str).str.startswith(str(phone)))].to_dict(orient='records')

def time_difference(start_time, end_time = None):
    if end_time is None or end_time == '':
        end_time = datetime.datetime.now() # + datetime.timedelta(hours=5, minutes=30)
    else:
        end_time = datetime.datetime.strptime(end_time, "%Y-%m-%dT%H:%M")
    if "T" in start_time:
        start_time = datetime.datetime.strptime(start_time, "%Y-%m-%dT%H:%M")
    else:
        start_time = datetime.datetime.strptime(start_time, "%d-%m-%Y %I:%M %p")
    diff = end_time - start_time
    return diff.days, (diff.seconds//3600)%24

def time_difference_str(start_time, end_time = None):
    days, hrs = time_difference(start_time, end_time)
    return str(days) + ' Days ' + str(hrs) + ' Hrs'

def fetch_payments(rid=None, page=None, count=None, date=None, **kwargs):
    register_df = read_csv(REGISTER_DB)
    customer_df = read_csv(CUSTOMERS_DB)
    transaction_df = read_csv(TRANSACTIONS_DB)

    if rid is not None:
        transaction_df['rid'] = transaction_df['rid'].astype(type(rid))
        transaction_df = transaction_df[transaction_df['rid'] == rid]
    
    if page is not None and count is not None:
        page, count = int(page), int(count)
        transaction_df['tt'] = pd.to_datetime(transaction_df['t'])
        transaction_df = transaction_df.sort_values('tt', ascending=False)
        transaction_df = transaction_df[page * count:(page + 1) * count]
    
    if date is not None:
        transaction_df = transaction_df[transaction_df['t'].str.startswith(date)]

    merged_transaction_register_df = transaction_df.merge(right=register_df, how='left', left_on='rid', right_on='id')
    merged_required = merged_transaction_register_df.merge(right=customer_df, how='left', left_on='cid', right_on='id')
    final_df = merged_required[['id_x', 'rno', 'n', 't', 'm', 'a_x', 'd', 'rid']]
    final_df = final_df.sort_values('t', ascending=False)
    final_df['n'] = final_df['n'].fillna('Credit')
    final_df['rno'] = final_df['rno'].fillna(-1)
    final_df['rno'] = final_df['rno'].astype(int)
    final_df = final_df.fillna('')
    final_df['t'] = final_df['t'].apply(transform_time)
    try:
        if kwargs.get('get_sum', False):
            return final_df['a_x'].sum()
    except Exception as err:
        return 'Unable to fetch amount ' + str(err)
    return final_df.to_dict(orient='records')

def transform_time(row):
    try:
        dt = datetime.datetime.strptime(row, '%Y-%m-%dT%H:%M')
        return dt.strftime('%d-%m-%Y %I:%M %p')
    except Exception as err:
        return None

def create_report(date):

    final_df = get_register_detail(date=date)

    year, month, date = date.split('-')

    docx = Document()
    header1 = docx.add_heading(HOTEL_NAME)
    header1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.style.font.size = Pt(36) # type: ignore
    header2 = docx.add_heading(HOTEL_ADDRESS, level=3)
    header2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.style.font.size = Pt(24) # type: ignore
    docx.add_paragraph('Date: ' + date + '-' + month + '-' + year).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = docx.add_table(rows=len(final_df) + 1, cols=6)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Sl. No.'
    table.rows[0].cells[1].text = 'Name & Address'
    table.rows[0].cells[2].text = 'Check In'
    table.rows[0].cells[3].text = 'Phone'
    table.rows[0].cells[4].text = 'Document'
    table.rows[0].cells[5].text = 'Purpose Of Visit'
    for idx, row in final_df.iterrows():
        idx_int = int(idx) # type: ignore
        table.rows[idx_int + 1].cells[0].text = str(idx_int + 1) # pyright: ignore[reportOperatorIssue]
        table.rows[idx_int + 1].cells[1].text = row['n'] + '\n' + row['a']
        table.rows[idx_int + 1].cells[2].text = row['cin']
        table.rows[idx_int + 1].cells[3].text = str(row['p'])
        table.rows[idx_int + 1].cells[4].text = row['it'] + '\n' + str(row['ip'])
        table.rows[idx_int + 1].cells[5].text = row['pov']
        table.rows[idx_int + 1].cells[0].width = Inches(0.5)
        table.rows[idx_int + 1].cells[1].width = Inches(3)
        table.rows[idx_int + 1].cells[2].width = Inches(2.5)
        table.rows[idx_int + 1].cells[3].width = Inches(0.5)
        table.rows[idx_int + 1].cells[4].width = Inches(2)
        table.rows[idx_int + 1].cells[5].width = Inches(2)
    if not os.path.exists(REPORT_PATH):
        os.mkdir(REPORT_PATH)
    filename = REPORT_PATH
    if REPORT_MONTHLY_FLODER:
        str_month = {
            1: 'January',
            2: 'February',
            3: 'March',
            4: 'April',
            5: 'May',
            6: 'June',
            7: 'July',
            8: 'August',
            9: 'September',
            10: 'October',
            11: 'November',
            12: 'December'
        }
        cnt_str_month = str_month[int(month)]
        month_floder = os.path.join(REPORT_PATH, cnt_str_month + ' ' + year)
        if not os.path.exists(month_floder):
            os.mkdir(month_floder)
        filename = month_floder
    os.startfile(filename)
    docx.save(os.path.join(filename, date + '-' + month + '-' + year + '.docx'))
    return os.path.join(filename, date + '-' + month + '-' + year + '.docx')

def get_report_records(month, year):
    day = 1
    df = read_csv(REGISTER_DB)
    res = {}
    while True:
        try:
            datetime.datetime(month=int(month), year=int(year), day=day)
            date = year + '-' + month.zfill(2) + '-' + str(day).zfill(2)
            count = len(df[df['cin'].str.startswith(date)])
            res[date] = {'nor': count}
            day += 1
            if day > 31:
                break
        except Exception as err:
            break
    return res

def _send_webhook_alert(data):
    time.sleep(3)
    res = requests.post(DISCORD_WEBHOOK_URL, json={'embeds': [data]})
    if 200 <= res.status_code < 300:
        return {'success': True, 'message': 'alerted successfully!'}
    else:
        with open(ALERT_DB, 'r') as f:
            f_data = json.load(f)
        f_data.append(data)
        with open(ALERT_DB, 'w') as f:
            json.dump(f_data, f)
        print(str(res.status_code) + ' ' + res.text)
        return {'success': False, 'message': 'Unable to send! Please try again later'}

def send_webhook_alert(data):
    threading.Thread(target=_send_webhook_alert, args=(data, )).start()

def push_webhook_alerts():
    f = open(ALERT_DB, 'r')
    try:
        f_data = json.load(f)
    except Exception as exe:
        send_webhook_alert({'Title': 'alert file format error', 'description': str(exe)})
        return
    f = open(ALERT_DB, 'w')
    f.write('[]')
    f.close()
    for data in f_data:
        time.sleep(2)
        send_webhook_alert(data)
    return None

def editDistance(s1, s2):
    m = len(s1)
    n = len(s2)
    prev = [0] * (n + 1)
    curr = [0] * (n + 1)
    for j in range(n + 1):
        prev[j] = j
    for i in range(1, m + 1):
        curr[0] = i  # j = 0
        for j in range(1, n + 1):
            if s1[i - 1] == s2[j - 1]:
                curr[j] = prev[j - 1]
            else:
                curr[j] = 1 + min(curr[j - 1], prev[j], prev[j - 1])
        prev = curr[:]    
    return prev[n]

if __name__ == "__main__":
    default_values()
